import streamlit as st
import openai
from docx import Document
from docx.shared import Inches
import io, json, base64, re, cv2, tempfile
from streamlit_drawable_canvas import st_canvas
from PIL import Image

st.set_page_config(page_title="HK Pro v9 - Video & Word", layout="wide")

# --- 1. KONFIGURACJA ---
if "OPENAI_API_KEY" in st.secrets:
    client = openai.OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
else:
    st.error("Błąd: Skonfiguruj klucz API!")
    st.stop()

if 'step' not in st.session_state: st.session_state.step = 1
if 'data' not in st.session_state: st.session_state.data = {}

# --- 2. FUNKCJE TECHNICZNE (OBRAZ I WIDEO) ---
def compress_img(image):
    if image.mode != "RGB": image = image.convert("RGB")
    image.thumbnail((1000, 1000))
    buf = io.BytesIO()
    image.save(buf, format="JPEG", quality=80)
    return base64.b64encode(buf.getvalue()).decode('utf-8')

def process_video(video_file, max_frames=15):
    """Wyciąga klatki z wideo dla AI"""
    tfile = tempfile.NamedTemporaryFile(delete=False)
    tfile.write(video_file.read())
    cap = cv2.VideoCapture(tfile.name)
    frames = []
    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    step = max(1, total_frames // max_frames)
    
    count = 0
    while cap.isOpened() and len(frames) < max_frames:
        ret, frame = cap.read()
        if not ret: break
        if count % step == 0:
            # Konwersja BGR na RGB
            frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            img = Image.fromarray(frame_rgb)
            frames.append(compress_img(img))
        count += 1
    cap.release()
    return frames

# --- 3. FUNKCJE WORD ---
def merge_broken_tags(paragraphs):
    for p in paragraphs:
        full_text = "".join(run.text for run in p.runs)
        if "{{" in full_text and "}}" in full_text:
            for run in p.runs: run.text = ""
            p.runs[0].text = full_text

def get_tags(docx_bytes):
    doc = Document(io.BytesIO(docx_bytes))
    merge_broken_tags(doc.paragraphs)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells: merge_broken_tags(cell.paragraphs if hasattr(c, 'paragraphs') else [])
    
    full_text = " ".join([p.text for p in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells: full_text += " " + cell.text
            
    found = re.findall(r"\{\{(.*?)\}\}", full_text)
    cleaned = [t.strip() for t in found if t.strip()]
    text_tags = [t for t in set(cleaned) if 'podpis' not in t.lower()]
    sig_tags = [t for t in set(cleaned) if 'podpis' in t.lower()]
    return text_tags, sig_tags

def apply_to_doc(doc, text_map, sig_map):
    all_paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells: all_paras.extend(cell.paragraphs)
    
    merge_broken_tags(all_paras)

    for p in all_paras:
        # Podmiana podpisów
        for sig_tag, canv in sig_map.items():
            pattern = re.compile(r"\{\{\s*" + re.escape(sig_tag) + r"\s*\}\}", re.IGNORECASE)
            if pattern.search(p.text) and canv.image_data is not None and canv.image_data.any():
                p.text = pattern.sub("", p.text)
                img = Image.fromarray(canv.image_data.astype('uint8'), 'RGBA')
                b = io.BytesIO(); img.save(b, format='PNG'); b.seek(0)
                p.add_run().add_picture(b, width=Inches(1.8))
        
        # Podmiana tekstu
        for txt_tag, val in text_map.items():
            pattern = re.compile(r"\{\{\s*" + re.escape(txt_tag) + r"\s*\}\}", re.IGNORECASE)
            if pattern.search(p.text):
                p.text = pattern.sub(str(val), p.text)

# --- 4. INTERFEJS ---
st.title("📄 Home Keeper Pro v9 - Zdjęcia + Wideo")

if st.session_state.step == 1:
    c1, c2 = st.columns(2)
    with c1: uploaded_word = st.file_uploader("1. Wgraj wzór Word (.docx)", type="docx")
    with c2: uploaded_files = st.file_uploader("2. Wgraj ZDJĘCIA i FILMY", type=["jpg", "png", "jpeg", "mp4", "mov"], accept_multiple_files=True)

    if st.button("🚀 ANALIZUJ MULTIMEDIA"):
        if uploaded_word and uploaded_files:
            with st.spinner("AI analizuje zdjęcia i klatki z filmu..."):
                w_bytes = uploaded_word.read()
                t_tags, s_tags = get_tags(w_bytes)
                
                all_base64_images = []
                for f in uploaded_files:
                    if f.type.startswith("image"):
                        all_base64_images.append(compress_img(Image.open(f)))
                    elif f.type.startswith("video") or f.name.endswith(('.mp4', '.mov')):
                        all_base64_images.extend(process_video(f))
                
                # Ograniczamy do 20 obrazów, by nie przekroczyć limitów API
                final_payload = all_base64_images[:20]
                
                prompt = f"""
                Jesteś ekspertem ds. inwentaryzacji. Na podstawie zdjęć i klatek z filmu wypełnij pola: {t_tags}.
                ZASADY: 
                - WYPISZ KAŻDY SZCZEGÓŁ OSOBNO (lista po przecinku).
                - ZAKAZ STRESZCZANIA (nie pisz 'klucze', pisz '1x piwnica, 2x wejściowe').
                - Jeśli na filmie/zdjęciu są kody lub liczby, podaj je precyzyjnie.
                Zwróć TYLKO JSON: {{"tag": "wartość"}}
                """
                
                res = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[{"role":"user","content":[{"type":"text","text":prompt},
                    *[{"type":"image_url","image_url":{"url":f"data:image/jpeg;base64,{i}"}} for i in final_payload]]}],
                    response_format={"type": "json_object"}
                )
                
                st.session_state.data = json.loads(res.choices[0].message.content)
                st.session_state.t_tags, st.session_state.s_tags = t_tags, s_tags
                st.session_state.template = w_bytes
                st.session_state.step = 2
                st.rerun()

elif st.session_state.step == 2:
    st.subheader("📝 Weryfikacja danych i Podpisy")
    col_text, col_sig = st.columns([1, 1])
    
    with col_text:
        st.write("🔍 **Zweryfikuj treść:**")
        updated_text = {tag: st.text_area(f"Pole: {tag}", value=st.session_state.data.get(tag, ""), height=100) for tag in st.session_state.t_tags}
    
    with col_sig:
        st.write("🖋️ **Złóż podpisy:**")
        final_sigs = {tag: st_canvas(height=150, width=400, key=f"sig_{tag}", background_color="#f5f5f5", display_toolbar=False) for tag in st.session_state.s_tags}
        for tag in st.session_state.s_tags: st.caption(f"Miejsce na: {tag}")

    if st.button("🖨️ GENERUJ RAPORT WORD"):
        doc = Document(io.BytesIO(st.session_state.template))
        apply_final_changes = apply_to_doc(doc, updated_text, final_sigs)
        out = io.BytesIO(); doc.save(out)
        st.success("✅ Gotowe!")
        st.download_button("📥 POBIERZ RAPORT (.docx)", out.getvalue(), "raport_hk.docx")

    if st.button("⬅️ Zacznij od nowa"):
        st.session_state.step = 1; st.rerun()
